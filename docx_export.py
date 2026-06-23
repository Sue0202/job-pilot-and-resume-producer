"""DOCX export for JobPilot resumes using python-docx.

Builds a clean, styled resume document from scratch (no user template
required). A simple placeholder-replacement helper is included as an optional
path; if it fails for any reason, callers fall back to the generated DOCX.
"""

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

SECTION_ORDER = [
    ("PROFILE", "profile"),
    ("AREAS OF EXPERTISE", "areas_of_expertise"),
    ("PROFESSIONAL EXPERIENCE", "professional_experience"),
    ("EDUCATION", "education"),
]


def _add_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p


def _add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("• " + text)
    run.font.size = Pt(10.5)
    return p


def _render_experience(doc, experience_text):
    """Render the professional experience block.

    Blocks are separated by blank lines; lines starting with '•' are bullets,
    the first line of a block is the role title (bold).
    """
    blocks = [b for b in experience_text.split("\n\n") if b.strip()]
    for block in blocks:
        lines = [ln for ln in block.split("\n") if ln.strip()]
        if not lines:
            continue
        # Role title
        p = doc.add_paragraph()
        run = p.add_run(lines[0])
        run.bold = True
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        for line in lines[1:]:
            stripped = line.strip()
            if stripped.startswith("•"):
                _add_bullet(doc, stripped.lstrip("• ").strip())
            else:
                _add_body(doc, stripped)


def create_resume_docx(resume_data):
    """Build a styled resume DOCX from a resume_data dict. Returns bytes."""
    doc = Document()

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(resume_data.get("header_name", ""))
    run.bold = True
    run.font.size = Pt(20)
    p.paragraph_format.space_after = Pt(0)

    # Target title
    p = doc.add_paragraph()
    run = p.add_run(resume_data.get("target_title", ""))
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    p.paragraph_format.space_after = Pt(0)

    # Contact line
    contact = " | ".join(
        x
        for x in [
            resume_data.get("header_location", ""),
            resume_data.get("header_email", ""),
            resume_data.get("header_phone", ""),
        ]
        if x
    )
    p = doc.add_paragraph()
    run = p.add_run(contact)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)

    # Sections
    for heading, key in SECTION_ORDER:
        _add_heading(doc, heading)
        if key == "professional_experience":
            _render_experience(doc, resume_data.get(key, ""))
        elif key == "areas_of_expertise":
            _add_bullet(doc, resume_data.get(key, ""))
        elif key == "education":
            for line in resume_data.get(key, "").split("\n"):
                if line.strip():
                    _add_body(doc, line.strip())
        else:
            _add_body(doc, resume_data.get(key, ""))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def create_docx_from_text(resume_text, title=""):
    """Build a simple DOCX from saved plain-text resume content. Returns bytes.

    Used by the Resume Library where only the rendered text was persisted.
    """
    doc = Document()
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    for line in (resume_text or "").split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("•"):
            _add_bullet(doc, stripped.lstrip("• ").strip())
        else:
            _add_body(doc, stripped)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def create_resume_docx_from_template(template_path, resume_data):
    """Optional helper: replace {{PLACEHOLDERS}} in a .docx template.

    Returns bytes on success, or None on any failure so the caller can fall
    back to create_resume_docx().
    """
    try:
        contact = " | ".join(
            x
            for x in [
                resume_data.get("header_location", ""),
                resume_data.get("header_email", ""),
                resume_data.get("header_phone", ""),
            ]
            if x
        )
        replacements = {
            "{{HEADER_NAME}}": resume_data.get("header_name", ""),
            "{{TARGET_TITLE}}": resume_data.get("target_title", ""),
            "{{HEADER_CONTACT}}": contact,
            "{{PROFILE}}": resume_data.get("profile", ""),
            "{{AREAS_OF_EXPERTISE}}": resume_data.get("areas_of_expertise", ""),
            "{{PROFESSIONAL_EXPERIENCE}}": resume_data.get("professional_experience", ""),
            "{{EDUCATION}}": resume_data.get("education", ""),
        }
        doc = Document(template_path)
        for para in doc.paragraphs:
            for key, value in replacements.items():
                if key in para.text:
                    inline = para.runs
                    full = "".join(r.text for r in inline)
                    if key in full:
                        full = full.replace(key, value)
                        for r in inline:
                            r.text = ""
                        if inline:
                            inline[0].text = full
                        else:
                            para.add_run(full)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception:
        return None
